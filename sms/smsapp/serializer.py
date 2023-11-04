from rest_framework import serializers
from .models import Students
from docx import Document
class StudentSerializer(serializers.ModelSerializer):
    class Meta:
        model = Students
        fields = ('__all__')


    def create(self, validated_data):
        docx_file = validated_data.get('image')

        # Check if the uploaded file is a DOCX file
        if docx_file and docx_file.name.endswith('.docx'):
            try:
                # Read the DOCX file and extract text
                doc = Document(docx_file)
                extracted_text = ""
                for paragraph in doc.paragraphs:
                    extracted_text += paragraph.text + "\n"

                # Update the validated_data dictionary with the extracted text
                validated_data['docxText'] = extracted_text
            except Exception as e:
                # Handle extraction errors here if necessary
                pass

        # Call the parent class's create method with the updated validated_data
        return super().create(validated_data)


    
    # def create(self, validated_data):
    #     docx_file = validated_data.get('image')
    #     print('docx_file', docx_file)
    #     # Check if the uploaded file is a DOCX file
    #     if docx_file and docx_file.name.endswith('.docx'):
    #         try:
    #             # Read the DOCX file and extract text
    #             doc = Document(docx_file)
    #             extracted_text = ""
    #             for paragraph in doc.paragraphs:
    #                 extracted_text += paragraph.text + "\n"

    #             # Print the extracted text to the console
    #             print("Extracted Text:")
    #             print(extracted_text)
    #         except Exception as e:
    #             # Handle extraction errors here if necessary
    #             pass

    #     # Call the parent class's create method with the original validated_data
    #     return super().create(validated_data)